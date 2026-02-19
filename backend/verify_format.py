
import os
import sys
import django
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Setup Django
# Since we will run from backend/, the current dir is backend/
# which contains api/, config/ etc.
sys.path.append(os.getcwd())
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "config.settings")
django.setup()

from api.services import save_improved_document

# Sample text with markdown
sample_text = """# Заголовок 1
## Заголовок 2
Обычный текст параграфа. Он должен быть выровнен по ширине.
А это **жирный текст** внутри параграфа.
"""

# Generate
print("Generating document...")
content_file = save_improved_document(sample_text, "test_formatting.docx")

# Save to disk to check
with open("test_formatting_result.docx", "wb") as f:
    f.write(content_file.read())

print("Document saved to test_formatting_result.docx")

# Inspect
doc = Document("test_formatting_result.docx")

print("--- Inspection ---")
# Para 1: Heading 1
try:
    p1 = doc.paragraphs[0]
    print(f"Para 1 style: {p1.style.name} (Expected: Heading 1)")
    print(f"Para 1 text: {p1.text}")

    # Para 3: Normal text
    p3 = doc.paragraphs[2]
    print(f"Para 3 text: {p3.text}")
    print(f"Para 3 alignment: {p3.alignment} (Expected: {WD_ALIGN_PARAGRAPH.JUSTIFY})")

    # Check runs for font
    if p3.runs:
        run = p3.runs[0]
        print(f"Run font: {run.font.name} (Expected: Times New Roman)")
        print(f"Run size: {run.font.size} (Expected: {Pt(12)})")
    else:
        print("Para 3 has no runs?")

    # Para 4: Bold check
    p4 = doc.paragraphs[3]
    print(f"Para 4 text: {p4.text}")
    # "А это ", "жирный текст", " внутри параграфа."
    if len(p4.runs) > 1:
        print(f"Run 1 bold: {p4.runs[1].bold} (Expected: True)")
        print(f"Run 1 text: {p4.runs[1].text}")
    else:
        print("Para 4 has unexpected run count.")

    print("--- End Inspection ---")
except Exception as e:
    print(f"Inspection error: {e}")
