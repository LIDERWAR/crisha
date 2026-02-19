import fitz  # PyMuPDF
import openai
import os
import json
from django.conf import settings
import docx
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage

# Инициализация клиента OpenAI (совместимый с DeepSeek)
client = openai.OpenAI(
    api_key=os.getenv("DEEPSEEK_API_KEY"),
    base_url="https://api.polza.ai/api/v1"
)

def extract_text_from_pdf(file_stream):
    """
    Извлекает текст из потока PDF-файла, используя PyMuPDF.
    """
    try:
        doc = fitz.open(stream=file_stream.read(), filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        return text
    except Exception as e:
        print(f"Ошибка извлечения текста (PDF): {e}")
        return None

def extract_text_from_docx(file_stream):
    """
    Извлекает текст из .docx файла.
    """
    try:
        doc = docx.Document(file_stream)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        print(f"Ошибка извлечения текста (DOCX): {e}")
        return None

def extract_text_from_txt(file_stream):
    """
    Извлекает текст из .txt файла.
    """
    try:
        return file_stream.read().decode('utf-8')
    except Exception as e:
        return None

def convert_doc_to_docx(doc_path):
    """
    Конвертирует .doc в .docx используя MS Word через COM интерфейс.
    Возвращает путь к новому файлу .docx.
    ВНИМАНИЕ: Требует установленного MS Word на сервере/машине.
    """
    if not os.path.exists(doc_path):
        print(f"Файл не найден для конвертации: {doc_path}")
        return None

    try:
        import win32com.client
        import pythoncom
        
        # Инициализация COM в текущем потоке
        pythoncom.CoInitialize()
        
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        # Открываем .doc
        # Absolute path is valid, doc_path comes from django which usually is absolute if file system storage
        abs_path = os.path.abspath(doc_path)
        doc = word.Documents.Open(abs_path)
        
        # Путь для .docx
        docx_path = abs_path + "x" # .doc -> .docx
        
        # Сохраняем как .docx (FileFormat=16 for wdFormatXMLDocument)
        doc.SaveAs(docx_path, FileFormat=16)
        doc.Close()
        # word.Quit() # Не закрываем Word полностью, чтобы не тормозить, или закрываем? 
        # Лучше закрыть Quit, если мы не хотим держать процесс. 
        # Но если много запросов, это медленно. Для прототипа Quit ок.
        word.Quit()
        
        return docx_path
    except Exception as e:
        print(f"Ошибка конвертации DOC -> DOCX: {e}")
        # Пытаемся закрыть Word если ошибка
        try:
            if 'word' in locals():
                word.Quit()
        except:
            pass
        return None

def analyze_contract_with_ai(contract_text):
    """
    Отправляет текст договора в DeepSeek для юридического анализа.
    Возвращает JSON с рисками, рекомендациями и УЛУЧШЕННЫМ текстом.
    """
    api_key = os.getenv("DEEPSEEK_API_KEY")
    use_mock = os.getenv("USE_MOCK_AI", "False").lower() == "true"
    
    # Проверка на заглушку
    # Если стоит USE_MOCK_AI=True, то всегда мок
    # Если ключа нет, то тоже мок
    print(f"DEBUG: api_key exists: {bool(api_key)}, startswith sk-placeholder: {api_key.startswith('sk-placeholder') if api_key else 'N/A'}, use_mock: {use_mock}")
    if use_mock or (not api_key or api_key.startswith("sk-placeholder")):
        print("--- ИСПОЛЬЗУЮ MOCK AI (Режим симуляции или нет ключа) ---")
        
        # Генерируем "улучшенный" текст на основе исходного
        # В реальной жизни тут AI переписывает. В эмуляторе просто делаем замену.
        improved_text = contract_text.replace("1%", "0.1% (исправлено AI)").replace("одностороннем порядке", "судебном порядке (исправлено AI)")
        if improved_text == contract_text:
             improved_text = contract_text + "\n\n[ДОПОЛНЕНИЕ ОТ AI: В тексте не найдено явных критических угроз для автозамены в режиме симуляции, но текст проверен.]"
        
        return {
            "score": 85,
            "summary": "Это тестовый анализ (Режим симуляции). Договор проверен, критичные места отмечены.",
            "risks": [
                {"title": "Право одностороннего расторжения", "description": "Арендодатель может расторгнуть договор без суда с уведомлением за 30 дней.", "severity": "high"},
                {"title": "Штрафы за просрочку", "description": "Пеня составляет 1% от суммы за каждый день просрочки.", "severity": "low"}
            ],
            "recommendations": [
                {
                    "title": "Предложить протокол разногласий", 
                    "description": "Необходимо исключить право одностороннего расторжения. Это создает риск внезапного выселения.",
                    "clause_example": "Арендодатель имеет право расторгнуть договор только в судебном порядке при существенном нарушении условий Арендатором."
                },
                {
                    "title": "Снизить пеню", 
                    "description": "Пеня 1% в день - это 365% годовых. Это кабальная сделка. Нормальная практика - 0.1%.",
                    "clause_example": "За просрочку платежа Арендатор уплачивает пени в размере 0.1% от суммы задолженности за каждый день просрочки."
                }
            ],
            "rewritten_text": improved_text
        }

    if not contract_text or len(contract_text) < 50:
        return {"error": "Текст слишком короткий или пустой."}

    prompt = (
        "Ты профессиональный юрист. Проанализируй договор. "
        "Твоя главная задача: найти риски и ПЕРЕПИСАТЬ ВЕСЬ ДОГОВОР ЦЕЛИКОМ, устранив эти риски, но сохранив остальной текст и структуру."
        "\\n\\n"
        "Формат JSON: "
        "{"
        "  'score': (0-100), "
        "  'summary': '...', "
        "  'risks': [{'title': '...', 'description': '...', 'severity': '...'}], "
        "  'recommendations': [{'title': '...', 'description': '...', 'clause_example': '...'}], "
        "  'rewritten_text': 'ПОЛНЫЙ текст документа от начала до конца. Если ты меняешь пункт, меняй его в тексте. Если пункт нормальный - оставляй как есть. Текст должен быть готов к подписанию.' "
        "}"
        "\\n\\n"
        f"Текст договора:\\n{contract_text[:14000]}" 
    )

    try:
        print("--- Отправка запроса к DeepSeek ---")
        response = client.chat.completions.create(
            model="deepseek/deepseek-chat",
            messages=[
                {"role": "system", "content": "Output valid JSON only."},
                {"role": "user", "content": prompt}
            ],
            response_format={"type": "json_object"}
        )
        content = response.choices[0].message.content
        
        if content.startswith("```"):
            content = content.strip("`").replace("json\n", "").replace("json", "")
        
        result = json.loads(content)
        if 'score' in result: result['score'] = int(result['score'])
             
        return result
    except Exception as e:
        print(f"!!! Ошибка сервиса AI: {e} !!!")
        return {
            "score": 0,
            "summary": f"Ошибка AI: {str(e)}",
            "risks": [],
            "recommendations": []
        }

def save_improved_document(text, original_filename):
    """
    Сохраняет улучшенный текст в файл (.docx или .txt) и возвращает объект ContentFile.
    """
    filename_base = os.path.splitext(original_filename)[0]
    
    # Всегда стараемся сохранить как DOCX для удобства
    try:
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = docx.Document()
        
        # Настройка стиля Normal
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0)
        
        # Разбор текста по параграфам
        paragraphs = text.split('\n')
        
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
                
            # Простейший парсинг Markdown
            # Заголовки
            if para_text.startswith('# '):
                p = doc.add_heading(para_text[2:], level=1)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                continue
            elif para_text.startswith('## '):
                p = doc.add_heading(para_text[3:], level=2)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                continue
            elif para_text.startswith('### '):
                p = doc.add_heading(para_text[4:], level=3)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                continue
                
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Жирный текст (**text**)
            parts = para_text.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                # Если индекс нечетный, значит это текст ВНУТРИ ** **, делаем жирным
                if i % 2 == 1:
                    run.bold = True

        # Сохраняем во временный буфер или сразу в storage
        # Django FileField принимает ContentFile
        from io import BytesIO
        f = BytesIO()
        doc.save(f)
        f.seek(0)
        return ContentFile(f.read(), name=f"{filename_base}_improved.docx")
    except Exception as e:
        print(f"Ошибка создания DOCX: {e}. Сохраняем как TXT.")
        return ContentFile(text.encode('utf-8'), name=f"{filename_base}_improved.txt")
