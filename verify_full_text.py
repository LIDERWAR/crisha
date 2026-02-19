
import requests
import os
import docx

# Create a test DOCX
doc = docx.Document()
doc.add_paragraph("ДОГОВОР АРЕНДЫ")
doc.add_paragraph("1. Арендодатель имеет право расторгнуть договор в одностороннем порядке.")
doc.add_paragraph("2. Пеня составляет 1% в день.")
doc.add_paragraph("3. Этот текст должен сохраниться полностью.")
doc.save("test_full_text.docx")

# API Endpoint
url = "http://127.0.0.1:8000/api/analyze/"

# Send POST request
with open("test_full_text.docx", "rb") as f:
    files = {"file": f}
    try:
        response = requests.post(url, files=files)
        print(f"Status Code: {response.status_code}")
        if response.status_code == 201:
            data = response.json()
            if "improved_file" in data and data["improved_file"]:
                print(f"Improved file URL: {data['improved_file']}")
                # Download and check content
                full_url = "http://127.0.0.1:8000" + data['improved_file']
                r_file = requests.get(full_url)
                if r_file.status_code == 200:
                    with open("downloaded_improved.docx", "wb") as f_down:
                        f_down.write(r_file.content)
                    
                    # Read content
                    doc_down = docx.Document("downloaded_improved.docx")
                    full_text = "\n".join([p.text for p in doc_down.paragraphs])
                    print("--- DOWNLOADED TEXT ---")
                    print(full_text)
                    
                    if "Этот текст должен сохраниться полностью" in full_text and "0.1%" in full_text:
                        print("SUCCESS: Full text preserved and modifications applied.")
                    else:
                        print("FAILURE: Text content check failed.")
                        
                else:
                    print(f"FAILURE: Could not download file. Status: {r_file.status_code}")
            else:
                print("No improved_file found in response.")
        else:
            print("Error:", response.text)
    except Exception as e:
        print(f"Request failed: {e}")

# Cleanup
if os.path.exists("test_full_text.docx"):
    os.remove("test_full_text.docx")
if os.path.exists("downloaded_improved.docx"):
    os.remove("downloaded_improved.docx")
